import os, json, boto3
from datetime import datetime
from docx import Document
from docx.shared import Pt

# ENV
TABLE_NAME = os.environ["WORKSHEETS_TABLE"]
RAW_BUCKET = os.environ["RAW_BUCKET"]
GENERATED_BUCKET = os.environ["GENERATED_BUCKET"]
GENERATED_PREFIX = os.environ.get("GENERATED_PREFIX", "generated-worksheets/")
BEDROCK_MODEL_ID = os.environ["BEDROCK_MODEL_ID"]

# AWS clients
dynamodb = boto3.resource("dynamodb")
table = dynamodb.Table(TABLE_NAME)
s3 = boto3.client("s3")
bedrock = boto3.client("bedrock-runtime")

# ---------- Helpers ----------
def fetch_processed_text(item):
    if "processedFileS3Path" in item:  # from uploaded file
        key = item["processedFileS3Path"]
        resp = s3.get_object(Bucket=RAW_BUCKET, Key=key)
        return resp["Body"].read().decode("utf-8")
    elif "processedText" in item:  # pasted text
        return item["processedText"]
    elif item.get("linkedContentId") == "RAG":
        return "Retrieved NCERT context (demo placeholder)"  # TODO integrate RAG
    return ""

def build_prompt(item, source_text):
    template = item.get("template", {})
    template_lines = [f"- {v.get('count',0)} {k} ({v.get('marks',0)} marks)" for k,v in template.items()]
    return f"""
You are an NCERT-aligned worksheet generator.
Output ONLY valid JSON matching this schema:
{{"worksheetId":"", "title":"", "language":"", "subject":"", "chapter":"",
"questions":[{{"qId":"","type":"","text":"","options":[],"modelAnswer":"","marks":0,"rubric":"","difficulty":""}}],
"totalMarks":0}}

WorksheetId: {item['worksheetId']}
Language: {item.get('language')}
Subject: {item.get('subject')}
Chapter: {item.get('chapter')}
Topic: {item.get('topic')}
Difficulty: {item.get('difficulty')}
Template:
{chr(10).join(template_lines)}

Source:
{source_text}
"""

def call_bedrock(prompt):
    resp = bedrock.invoke_model(
        modelId=BEDROCK_MODEL_ID,
        contentType="application/json",
        accept="application/json",
        body=json.dumps({"input": prompt})
    )
    return resp["body"].read().decode("utf-8")

def extract_json(raw_text):
    start, end = raw_text.find("{"), raw_text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("No JSON found")
    return json.loads(raw_text[start:end+1])

def generate_docx(qjson, include_answers=False):
    doc = Document()
    doc.add_heading(qjson.get("title", "Worksheet"), 1)
    for q in qjson.get("questions", []):
        para = doc.add_paragraph(f"{q['qId']}. ({q['marks']}m) {q['text']}")
        para.runs[0].font.size = Pt(12)
        if q["type"].upper() == "MCQ":
            for opt in q.get("options", []):
                doc.add_paragraph(f"   {opt}")
        if include_answers and q.get("modelAnswer"):
            doc.add_paragraph(f"   [Answer: {q['modelAnswer']}]")
    import io
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def upload_to_s3(key, data, ctype):
    s3.put_object(Bucket=GENERATED_BUCKET, Key=key, Body=data, ContentType=ctype)
    return f"s3://{GENERATED_BUCKET}/{key}"

# ---------- Handler ----------
def lambda_handler(event, context):
    try:
        body = json.loads(event["body"]) if "body" in event else event
        worksheet_id = body.get("worksheetId")
        if not worksheet_id:
            return {"statusCode":400,"body":json.dumps({"error":"worksheetId required"})}

        # 1. Get worksheet
        resp = table.get_item(Key={"worksheetId": worksheet_id, "contentId":"NONE"})
        item = resp.get("Item")
        if not item:
            return {"statusCode":404,"body":json.dumps({"error":"Worksheet not found"})}

        # 2. Fetch text
        source_text = fetch_processed_text(item)

        # 3. Bedrock call
        prompt = build_prompt(item, source_text)
        raw = call_bedrock(prompt)
        qjson = extract_json(raw)
        qjson["worksheetId"] = worksheet_id
        qjson["generatedAt"] = datetime.utcnow().isoformat()

        # 4. Save JSON
        json_key = f"{GENERATED_PREFIX}{worksheet_id}.json"
        json_path = upload_to_s3(json_key, json.dumps(qjson).encode("utf-8"), "application/json")

        # 5. Generate DOCX
        st_key = f"{GENERATED_PREFIX}{worksheet_id}_student.docx"
        st_path = upload_to_s3(st_key, generate_docx(qjson, False),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        ans_key = f"{GENERATED_PREFIX}{worksheet_id}_answerkey.docx"
        ans_path = upload_to_s3(ans_key, generate_docx(qjson, True),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # 6. Update DynamoDB
        table.update_item(
            Key={"worksheetId": worksheet_id, "contentId":"NONE"},
            UpdateExpression="SET generatedJsonS3Path=:j, studentDocxS3Path=:s, answerKeyDocxS3Path=:a, #st=:st, completedAt=:c",
            ExpressionAttributeValues={
                ":j": json_path, ":s": st_path, ":a": ans_path,
                ":st":"done", ":c": datetime.utcnow().isoformat()
            },
            ExpressionAttributeNames={"#st":"status"}
        )

        return {"statusCode":200,"body":json.dumps({
            "message":"Worksheet generated",
            "worksheetId": worksheet_id,
            "jsonPath": json_path,
            "studentDocxPath": st_path,
            "answerDocxPath": ans_path,
            "status":"done"
        })}
    except Exception as e:
        return {"statusCode":500,"body":json.dumps({"error":str(e)})}
